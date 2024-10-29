from django.db import models
from django.contrib.auth import get_user_model
from django.core.exceptions import ValidationError
from django.db import transaction
import jsonschema

User = get_user_model()

class DocumentTemplate(models.Model):
    name = models.CharField(max_length=255)
    content = models.TextField()
    fields = models.JSONField(default=list)
    version = models.IntegerField(default=1, editable=False)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    owners = models.ManyToManyField(
        User,
        related_name='document_templates',
        through='DocumentTemplateOwnership'
    )
    
    # Field schema for validation
    FIELDS_SCHEMA = {
        "type": "array",
        "items": {
            "type": "object",
            "properties": {
                "name": {"type": "string"},
                "type": {"type": "string", "enum": ["text", "number", "date", "choice"]},
                "required": {"type": "boolean"},
                "options": {"type": "array", "items": {"type": "string"}}
            },
            "required": ["name", "type"]
        }
    }

    class Meta:
        ordering = ['-created_at', '-version']
        constraints = [
            models.UniqueConstraint(
                fields=['name', 'version'],
                name='unique_template_version'
            )
        ]

    def clean(self):
        try:
            jsonschema.validate(instance=self.fields, schema=self.FIELDS_SCHEMA)
        except jsonschema.exceptions.ValidationError as e:
            raise ValidationError(f"Invalid fields format: {str(e)}")

    @transaction.atomic
    def save(self, *args, **kwargs):
        if self.pk:
            # Create new version
            self.pk = None
            self.version += 1
        super().save(*args, **kwargs)

    def export_as_pdf(self):
        try:
            # PDF generation logic
            from reportlab.pdfgen import canvas
            from io import BytesIO
            
            buffer = BytesIO()
            p = canvas.Canvas(buffer)
            
            # Add content to PDF
            p.drawString(100, 800, self.name)
            y_position = 750
            for line in self.content.split('\n'):
                p.drawString(100, y_position, line)
                y_position -= 20
            
            p.showPage()
            p.save()
            
            return buffer.getvalue()
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"PDF generation failed: {str(e)}")
            return None

    def export_as_word(self):
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            doc.add_heading(self.name, 0)
            doc.add_paragraph(self.content)
            
            buffer = BytesIO()
            doc.save(buffer)
            
            return buffer.getvalue()
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Word document generation failed: {str(e)}")
            return None

class DocumentTemplateOwnership(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    template = models.ForeignKey(DocumentTemplate, on_delete=models.CASCADE)
    access_level = models.CharField(
        max_length=20,
        choices=[
            ('view', 'View Only'),
            ('edit', 'Edit'),
            ('admin', 'Admin'),
        ],
        default='view'
    )
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        unique_together = ['user', 'template']
